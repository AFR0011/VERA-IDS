#!/usr/bin/env python3
"""
20.BuildThesisFinalizationPack.py
=================================

Purpose
-------
Build the final thesis-only reporting layer from already-materialized canonical
artifacts. This script does not train models or rerun Protocol A/Protocol B.

Outputs
-------
- thesis_finalization_pack/tables/*.csv
- thesis_finalization_pack/figures/*.png
- thesis_finalization_pack/output_manifest.csv
- Thesis and Paper 1 Drafts/Seminar Presentation/Thesis_Manuscript_Final_v3.md
- Thesis and Paper 1 Drafts/Seminar Presentation/Thesis_Manuscript_Final_v3.docx
"""

from __future__ import annotations

import csv
import json
import math
import os
import re
from collections import defaultdict
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Iterable

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


from ids_eval_framework.src.paths import REPO_ROOT

ROOT = REPO_ROOT

CFG: dict[str, Any] = {
    "thesis_pack": ROOT / "thesis_full_scope_pack",
    "step5_root": ROOT / "protocolB_grid_runs step 5 - open-set baselines",
    "step6_root": ROOT / "protocolB_grid_runs step 6 - sink-aware rejection",
    "drift_action_root": ROOT / "drift_action_study",
    "processed_v5": ROOT / "processed_V5",
    "processed_cicids_recovery": ROOT / "processed_V5_cicids17_recovery",
    "support_audit_cicids_recovery": ROOT / "protocolB_support_audit_out_cicids17_recovery",
    "out_root": ROOT / "thesis_finalization_pack",
    "source_manuscript": ROOT / "Thesis and Paper 1 Drafts" / "Thesis_Manuscript_Complete_v1.md",
    "manuscript_dir": ROOT / "Thesis and Paper 1 Drafts" / "Seminar Presentation",
    "final_md_name": "Thesis_Manuscript_Final_v3.md",
    "final_docx_name": "Thesis_Manuscript_Final_v3.docx",
    "score_chunksize": 250_000,
    "calibration_bins": 10,
    "overconfidence_threshold": 0.80,
    "selected_curve_holdouts": [
        ("CICIDS2017", "DDoS"),
        ("CICIDS2017", "BruteForce"),
        ("CICIoT2023", "Botnet"),
        ("CICIoT2023", "Scan/Recon"),
    ],
}


TABLES: list[tuple[str, str, str]] = [
    ("T20", "drift_action_cross_dataset_summary.csv", "Drift-action cross-dataset maintenance summary"),
    ("T21", "per_family_calibration.csv", "Per-family Protocol B replay calibration"),
    ("T22", "sink_calibration_diagnostic.csv", "Sink-family calibration and overconfidence diagnostic"),
    ("T23", "holdout_difficulty_taxonomy.csv", "Holdout difficulty taxonomy"),
    ("T24", "rejection_tradeoff_summary.csv", "Rejection tradeoff summary"),
    ("T25", "runtime_resource_cost_summary.csv", "Runtime and resource-cost summary"),
    ("T26", "selected_statistical_summary.csv", "Selected main-body CIs and effect sizes"),
    ("T27", "representative_error_cases.csv", "Representative sink error cases"),
    ("T28", "support_leakage_audit_summary.csv", "Support and leakage audit summary"),
]

FIGURES: list[tuple[str, str, str]] = [
    ("F20", "threshold_sensitivity_selected_holdouts.png", "Threshold-sensitivity curves"),
    ("F21", "support_leakage_audit_visual.png", "Support/leakage-audit visual"),
    ("F22", "sink_to_example_bridge.png", "Aggregate sink to representative example bridge"),
    ("F23", "reliability_deployment_feasibility.png", "Reliability and deployment feasibility summary"),
]

INLINE_RE = re.compile(r"(\*\*.*?\*\*|`.*?`)")


@dataclass
class CalibrationStats:
    n: int = 0
    correct: int = 0
    conf_sum: float = 0.0
    true_prob_sum: float = 0.0
    true_prob_n: int = 0
    contexts: set[str] | None = None
    bin_count: np.ndarray | None = None
    bin_correct: np.ndarray | None = None
    bin_conf_sum: np.ndarray | None = None

    def ensure(self, n_bins: int) -> None:
        if self.contexts is None:
            self.contexts = set()
        if self.bin_count is None:
            self.bin_count = np.zeros(n_bins, dtype=np.int64)
            self.bin_correct = np.zeros(n_bins, dtype=np.int64)
            self.bin_conf_sum = np.zeros(n_bins, dtype=np.float64)


def mkdirs() -> None:
    for sub in ["tables", "figures"]:
        (CFG["out_root"] / sub).mkdir(parents=True, exist_ok=True)
    CFG["manuscript_dir"].mkdir(parents=True, exist_ok=True)


def read_csv(name: str) -> pd.DataFrame:
    path = CFG["thesis_pack"] / name
    if not path.exists():
        return pd.DataFrame()
    return pd.read_csv(path)


def read_json(path: Path) -> dict[str, Any]:
    if not path.exists():
        return {}
    with path.open("r", encoding="utf-8") as handle:
        return json.load(handle)


def write_table(df: pd.DataFrame, name: str) -> Path:
    path = CFG["out_root"] / "tables" / name
    path.parent.mkdir(parents=True, exist_ok=True)
    df.to_csv(path, index=False)
    return path


def savefig(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    plt.tight_layout()
    plt.savefig(path, dpi=220, bbox_inches="tight")
    plt.close()


def safe_case_name(holdout_family: str) -> str:
    return str(holdout_family).replace("/", "_").replace("\\", "_").replace(" ", "_")


def case_dir(dataset: str, holdout_family: str) -> Path:
    return CFG["step5_root"] / f"{dataset}__holdout_{safe_case_name(holdout_family)}__winner_replay"


def case_scores_path(dataset: str, holdout_family: str) -> Path:
    return case_dir(dataset, holdout_family) / "test_scores.csv.gz"


def selected_control_thresholds(dataset: str, holdout_family: str) -> tuple[float, float]:
    selected_path = case_dir(dataset, holdout_family) / "selected_methods.csv"
    selected = pd.read_csv(selected_path)
    row = selected.loc[(selected["split"] == "test") & (selected["method"] == "control_tau")]
    if row.empty:
        row = selected.loc[selected["method"] == "control_tau"]
    rec = row.iloc[0]
    params = json.loads(str(rec.get("selection_param_json", "{}")))
    tau = float(params.get("tau", rec.get("tau", 0.0)))
    return float(rec["thr_high"]), tau


def iter_score_chunks(dataset: str, holdout_family: str, usecols: list[str]) -> Iterable[pd.DataFrame]:
    path = case_scores_path(dataset, holdout_family)
    if not path.exists():
        return
    for chunk in pd.read_csv(path, usecols=usecols, chunksize=int(CFG["score_chunksize"])):
        yield chunk


def calibration_ece(stats: CalibrationStats) -> float:
    if not stats.n or stats.bin_count is None or stats.bin_correct is None or stats.bin_conf_sum is None:
        return float("nan")
    total = float(stats.n)
    ece = 0.0
    for count, correct, conf_sum in zip(stats.bin_count, stats.bin_correct, stats.bin_conf_sum):
        if int(count) == 0:
            continue
        acc = float(correct) / float(count)
        avg_conf = float(conf_sum) / float(count)
        ece += (float(count) / total) * abs(acc - avg_conf)
    return float(ece)


def build_per_family_calibration(b_best: pd.DataFrame) -> pd.DataFrame:
    n_bins = int(CFG["calibration_bins"])
    stats: dict[tuple[str, str], CalibrationStats] = defaultdict(CalibrationStats)
    usecols = [
        "y_stage1_attack",
        "y_stage2_family",
        "is_true_unknown",
        "fam_pred_family",
        "fam_pmax",
        "true_known_family_prob",
    ]
    for rec in b_best[["dataset", "holdout_family"]].drop_duplicates().to_dict("records"):
        dataset = str(rec["dataset"])
        holdout = str(rec["holdout_family"])
        for chunk in iter_score_chunks(dataset, holdout, usecols):
            known_attack = (
                pd.to_numeric(chunk["y_stage1_attack"], errors="coerce").fillna(0).astype(int).eq(1)
                & pd.to_numeric(chunk["is_true_unknown"], errors="coerce").fillna(0).astype(int).eq(0)
            )
            sub = chunk.loc[known_attack].copy()
            if sub.empty:
                continue
            sub["fam_pmax"] = pd.to_numeric(sub["fam_pmax"], errors="coerce")
            sub["correct"] = sub["fam_pred_family"].astype(str).eq(sub["y_stage2_family"].astype(str))
            sub["bin"] = np.minimum(
                np.floor(np.clip(sub["fam_pmax"].fillna(0.0).to_numpy(dtype=np.float64), 0.0, 1.0) * n_bins).astype(int),
                n_bins - 1,
            )
            sub["true_known_family_prob"] = pd.to_numeric(sub["true_known_family_prob"], errors="coerce")
            for fam, grp in sub.groupby("y_stage2_family"):
                key = (dataset, str(fam))
                st = stats[key]
                st.ensure(n_bins)
                st.contexts.add(holdout)
                st.n += int(len(grp))
                st.correct += int(grp["correct"].sum())
                st.conf_sum += float(grp["fam_pmax"].sum())
                true_prob = grp["true_known_family_prob"].dropna()
                st.true_prob_sum += float(true_prob.sum())
                st.true_prob_n += int(len(true_prob))
                for bin_idx, bin_grp in grp.groupby("bin"):
                    bi = int(bin_idx)
                    st.bin_count[bi] += int(len(bin_grp))
                    st.bin_correct[bi] += int(bin_grp["correct"].sum())
                    st.bin_conf_sum[bi] += float(bin_grp["fam_pmax"].sum())

    rows: list[dict[str, Any]] = []
    for (dataset, family), st in sorted(stats.items()):
        accuracy = float(st.correct / st.n) if st.n else float("nan")
        mean_confidence = float(st.conf_sum / st.n) if st.n else float("nan")
        rows.append(
            {
                "dataset": dataset,
                "family": family,
                "n_known_attack_rows": int(st.n),
                "n_holdout_contexts": len(st.contexts or set()),
                "holdout_contexts": "|".join(sorted(st.contexts or set())),
                "toplabel_accuracy": accuracy,
                "mean_toplabel_confidence": mean_confidence,
                "confidence_accuracy_gap": mean_confidence - accuracy,
                "ece_10bin": calibration_ece(st),
                "mean_true_family_probability": float(st.true_prob_sum / st.true_prob_n) if st.true_prob_n else float("nan"),
                "source": "step5 Protocol B winner replay test_scores known-attack rows",
            }
        )
    return pd.DataFrame(rows)


def predict_control_system(chunk: pd.DataFrame, thr_high: float, tau: float) -> tuple[np.ndarray, np.ndarray]:
    p_attack = pd.to_numeric(chunk["p_attack"], errors="coerce").fillna(0.0).to_numpy(dtype=np.float64)
    fam_pmax = pd.to_numeric(chunk["fam_pmax"], errors="coerce").fillna(0.0).to_numpy(dtype=np.float64)
    fam_pred = chunk["fam_pred_family"].astype(str).to_numpy(dtype=object)
    pred = np.where(p_attack < thr_high, "Benign", np.where(fam_pmax < tau, "Unknown", fam_pred))
    conf = np.where(pred == "Benign", 1.0 - p_attack, fam_pmax)
    conf = np.where(pred == "Unknown", 1.0 - fam_pmax, conf)
    return pred.astype(object), np.asarray(conf, dtype=np.float64)


def build_sink_calibration_and_examples(failure: pd.DataFrame, per_family: pd.DataFrame) -> tuple[pd.DataFrame, pd.DataFrame]:
    rows: list[dict[str, Any]] = []
    examples: list[dict[str, Any]] = []
    usecols = [
        "row_id",
        "y_stage1_attack",
        "y_stage2_family",
        "is_true_unknown",
        "p_attack",
        "fam_pred_family",
        "fam_pmax",
        "top2_margin",
        "stage2_entropy",
    ]
    for rec in failure.to_dict("records"):
        dataset = str(rec["dataset"])
        holdout = str(rec["holdout_family"])
        sink = str(rec["top_sink_family"])
        thr_high, tau = selected_control_thresholds(dataset, holdout)
        n_unknown = 0
        n_sink = 0
        n_high_sink = 0
        sink_conf_sum = 0.0
        unknown_conf_sum = 0.0
        best: dict[str, Any] | None = None
        best_conf = -1.0
        for chunk in iter_score_chunks(dataset, holdout, usecols):
            unk = pd.to_numeric(chunk["is_true_unknown"], errors="coerce").fillna(0).astype(int).eq(1)
            sub = chunk.loc[unk].copy()
            if sub.empty:
                continue
            pred, conf = predict_control_system(sub, thr_high, tau)
            sink_mask = pred == sink
            n_unknown += int(len(sub))
            unknown_conf_sum += float(np.sum(conf))
            n_sink += int(np.sum(sink_mask))
            if np.any(sink_mask):
                sink_conf = conf[sink_mask]
                n_high_sink += int(np.sum(sink_conf >= float(CFG["overconfidence_threshold"])))
                sink_conf_sum += float(np.sum(sink_conf))
                local_pos = int(np.argmax(sink_conf))
                if float(sink_conf[local_pos]) > best_conf:
                    best_conf = float(sink_conf[local_pos])
                    best_row = sub.loc[sink_mask].iloc[local_pos]
                    best = {
                        "dataset": dataset,
                        "holdout_family": holdout,
                        "row_id": int(best_row["row_id"]),
                        "true_family": holdout,
                        "predicted_sink_family": sink,
                        "system_prediction": sink,
                        "sink_confidence": best_conf,
                        "p_attack": float(best_row["p_attack"]),
                        "fam_pred_family": str(best_row["fam_pred_family"]),
                        "fam_pmax": float(best_row["fam_pmax"]),
                        "top2_margin": float(best_row["top2_margin"]),
                        "stage2_entropy": float(best_row["stage2_entropy"]),
                        "aggregate_top_sink_share": float(rec["top_sink_share"]),
                        "source_score_table": str(case_scores_path(dataset, holdout).relative_to(ROOT)),
                    }
        calib_row = per_family.loc[(per_family["dataset"] == dataset) & (per_family["family"] == sink)]
        sink_ece = float(calib_row["ece_10bin"].iloc[0]) if not calib_row.empty else float("nan")
        sink_gap = float(calib_row["confidence_accuracy_gap"].iloc[0]) if not calib_row.empty else float("nan")
        overconf_rate = float(n_high_sink / n_unknown) if n_unknown else float("nan")
        mean_sink_conf = float(sink_conf_sum / n_sink) if n_sink else float("nan")
        if overconf_rate >= 0.25:
            signal = "overconfident_unknown_absorption"
        elif (not math.isnan(sink_ece) and sink_ece >= 0.05) or (not math.isnan(sink_gap) and abs(sink_gap) >= 0.05):
            signal = "known_family_calibration_risk"
        else:
            signal = "low_or_targeted_calibration_signal"
        rows.append(
            {
                "dataset": dataset,
                "holdout_family": holdout,
                "top_sink_family": sink,
                "n_true_unknown": int(n_unknown),
                "top_sink_share_from_failure_table": float(rec["top_sink_share"]),
                "control_tau_top_sink_count": int(n_sink),
                "control_tau_top_sink_share": float(n_sink / n_unknown) if n_unknown else float("nan"),
                "mean_unknown_prediction_confidence": float(unknown_conf_sum / n_unknown) if n_unknown else float("nan"),
                "mean_sink_confidence": mean_sink_conf,
                "overconfident_sink_absorption_rate": overconf_rate,
                "sink_known_family_ece_10bin": sink_ece,
                "sink_known_family_confidence_accuracy_gap": sink_gap,
                "calibration_failure_signal": signal,
            }
        )
        if best is not None:
            examples.append(best)
    return pd.DataFrame(rows), pd.DataFrame(examples)


def build_holdout_taxonomy(
    b_best: pd.DataFrame,
    notes: pd.DataFrame,
    open_recs: pd.DataFrame,
    sink_recs: pd.DataFrame,
    sink_diag: pd.DataFrame,
) -> pd.DataFrame:
    rows: list[dict[str, Any]] = []
    base = b_best[["dataset", "holdout_family", "unknown_detection_rate", "false_unknown_rate_all_known", "overall_reject_rate", "macro_f1"]].copy()
    for rec in base.to_dict("records"):
        dataset = str(rec["dataset"])
        holdout = str(rec["holdout_family"])
        note_row = notes.loc[(notes["dataset"] == dataset) & (notes["holdout_family"] == holdout)]
        open_row = open_recs.loc[(open_recs["dataset"] == dataset) & (open_recs["holdout_family"] == holdout)]
        sink_row = sink_recs.loc[(sink_recs["dataset"] == dataset) & (sink_recs["holdout_family"] == holdout)]
        diag_row = sink_diag.loc[(sink_diag["dataset"] == dataset) & (sink_diag["holdout_family"] == holdout)]
        note = str(note_row["holdout_case_note"].iloc[0]) if not note_row.empty else ""
        open_decision = str(open_row["decision"].iloc[0]) if not open_row.empty else ""
        sink_decision = str(sink_row["decision"].iloc[0]) if not sink_row.empty else ""
        tags: list[str] = []
        if note == "strong case":
            category = "strong"
        elif sink_decision == "targeted_recovery_candidate":
            category = "targeted-recovery"
        elif sink_decision == "negative_tradeoff":
            category = "negative-tradeoff"
        elif note == "hard failure case":
            category = "hard-collapse"
        elif open_decision in {"macro_gain_with_udr_cost", "udr_gain_with_macro_cost"}:
            category = "costly-rejectable"
        else:
            category = "mixed"
        if note:
            tags.append(note.replace(" ", "-"))
        if open_decision:
            tags.append(open_decision.replace("_", "-"))
        if sink_decision and sink_decision != "neutral":
            tags.append(sink_decision.replace("_", "-"))
        if not diag_row.empty:
            tags.append(str(diag_row["calibration_failure_signal"].iloc[0]).replace("_", "-"))
        rows.append(
            {
                **rec,
                "difficulty_category": category,
                "secondary_tags": "|".join(tags),
                "open_set_review_decision": open_decision,
                "sink_aware_review_decision": sink_decision,
                "top_sink_family": str(diag_row["top_sink_family"].iloc[0]) if not diag_row.empty else "",
                "overconfident_sink_absorption_rate": float(diag_row["overconfident_sink_absorption_rate"].iloc[0]) if not diag_row.empty else float("nan"),
            }
        )
    return pd.DataFrame(rows).sort_values(["dataset", "holdout_family"]).reset_index(drop=True)


def build_rejection_tradeoff(open_recs: pd.DataFrame) -> pd.DataFrame:
    if open_recs.empty:
        return pd.DataFrame()
    out = open_recs.copy()
    out["macro_f1_cost_or_gain"] = np.select(
        [
            pd.to_numeric(out["delta_macro_f1_vs_control"], errors="coerce") > 0.005,
            pd.to_numeric(out["delta_macro_f1_vs_control"], errors="coerce") < -0.005,
        ],
        ["macro_gain", "macro_cost"],
        default="macro_neutral",
    )
    out["unknown_detection_cost_or_gain"] = np.select(
        [
            pd.to_numeric(out["delta_unknown_detection_vs_control"], errors="coerce") > 0.02,
            pd.to_numeric(out["delta_unknown_detection_vs_control"], errors="coerce") < -0.02,
        ],
        ["udr_gain", "udr_cost"],
        default="udr_neutral",
    )
    out["deployment_tradeoff_note"] = out["decision"].map(
        {
            "joint_improvement": "improves macro-F1 and unknown detection under review rules",
            "udr_gain_bounded_cost": "improves unknown detection with bounded operating cost",
            "udr_gain_with_macro_cost": "improves unknown detection but costs macro-F1",
            "macro_gain_with_udr_cost": "improves macro-F1 but weakens unknown detection",
            "neutral_or_mixed": "no clean deployment advantage over tau control",
        }
    ).fillna("review manually")
    cols = [
        "dataset",
        "holdout_family",
        "method",
        "decision",
        "delta_macro_f1_vs_control",
        "delta_unknown_detection_vs_control",
        "delta_false_unknown_rate_all_known",
        "delta_overall_reject_rate",
        "delta_macro_f1_ci_low",
        "delta_macro_f1_ci_high",
        "delta_unknown_detection_ci_low",
        "delta_unknown_detection_ci_high",
        "macro_f1_cost_or_gain",
        "unknown_detection_cost_or_gain",
        "deployment_tradeoff_note",
    ]
    return out[[c for c in cols if c in out.columns]].copy()


def root_stats(path: Path) -> dict[str, Any]:
    files = [p for p in path.rglob("*") if p.is_file()] if path.exists() else []
    if not files:
        return {
            "file_count": 0,
            "total_size_mb": 0.0,
            "earliest_file_mtime": "",
            "latest_file_mtime": "",
            "artifact_update_span_hours": float("nan"),
        }
    mtimes = [p.stat().st_mtime for p in files]
    return {
        "file_count": len(files),
        "total_size_mb": round(sum(p.stat().st_size for p in files) / (1024 * 1024), 3),
        "earliest_file_mtime": pd.Timestamp(min(mtimes), unit="s").isoformat(),
        "latest_file_mtime": pd.Timestamp(max(mtimes), unit="s").isoformat(),
        "artifact_update_span_hours": round((max(mtimes) - min(mtimes)) / 3600.0, 3),
    }


def build_runtime_summary() -> pd.DataFrame:
    lanes = [
        ("Protocol A core baseline", ROOT / "runs_two_stage_V5_A_core", "fresh closed-set core baseline"),
        ("Drift primary analysis", ROOT / "drift_primary_analysis", "windowed drift metrics"),
        ("XAI full scope", ROOT / "xai_full_scope", "SHAP/explanation outputs"),
        ("Step 5 open-set baselines", CFG["step5_root"], "full audited rejector frontier"),
        ("Step 6 sink-aware rejection", CFG["step6_root"], "full sink-aware replay"),
        ("Drift action study", CFG["drift_action_root"], "maintenance action replay"),
        ("Thesis finalization pack", CFG["out_root"], "reporting synthesis layer"),
    ]
    rows: list[dict[str, Any]] = []
    for lane, path, role in lanes:
        settings = {}
        settings_path = path / "execution_settings.json"
        if settings_path.exists():
            settings = read_json(settings_path)
        rows.append(
            {
                "lane": lane,
                "artifact_root": str(path.relative_to(ROOT)) if path.exists() else str(path),
                "role": role,
                **root_stats(path),
                "cpu_total": settings.get("cpu_total", ""),
                "case_parallel_workers": settings.get("case_parallel_workers", ""),
                "threads_per_worker": settings.get("threads_per_worker", ""),
                "use_xgb_gpu": settings.get("use_xgb_gpu", ""),
                "wallclock_evidence_status": "execution settings available; exact historical wall-clock not logged"
                if settings
                else "artifact file span only; not exact wall-clock",
            }
        )
    return pd.DataFrame(rows)


def build_selected_statistics(open_recs: pd.DataFrame, sink_recs: pd.DataFrame) -> pd.DataFrame:
    rows: list[dict[str, Any]] = []
    open_positive = open_recs.loc[
        open_recs["decision"].isin(["joint_improvement", "udr_gain_bounded_cost"])
    ].copy()
    for rec in open_positive.to_dict("records"):
        rows.append(
            {
                "result_family": "open_set_positive",
                "dataset": rec["dataset"],
                "holdout_family": rec["holdout_family"],
                "method_or_lane": rec["method"],
                "decision": rec["decision"],
                "delta_macro_f1": rec["delta_macro_f1_vs_control"],
                "delta_macro_f1_ci_low": rec["delta_macro_f1_ci_low"],
                "delta_macro_f1_ci_high": rec["delta_macro_f1_ci_high"],
                "delta_unknown_detection": rec["delta_unknown_detection_vs_control"],
                "delta_unknown_detection_ci_low": rec["delta_unknown_detection_ci_low"],
                "delta_unknown_detection_ci_high": rec["delta_unknown_detection_ci_high"],
            }
        )
    for rec in sink_recs.loc[
        sink_recs["decision"].isin(["targeted_recovery_candidate", "negative_tradeoff"])
    ].to_dict("records"):
        rows.append(
            {
                "result_family": "sink_aware_case_study",
                "dataset": rec["dataset"],
                "holdout_family": rec["holdout_family"],
                "method_or_lane": "sink_aware_reject",
                "decision": rec["decision"],
                "delta_macro_f1": rec["delta_macro_f1_mean"],
                "delta_macro_f1_ci_low": rec["delta_macro_f1_ci_low"],
                "delta_macro_f1_ci_high": rec["delta_macro_f1_ci_high"],
                "delta_unknown_detection": rec["delta_unknown_detection_mean"],
                "delta_unknown_detection_ci_low": rec["delta_unknown_detection_ci_low"],
                "delta_unknown_detection_ci_high": rec["delta_unknown_detection_ci_high"],
            }
        )
    return pd.DataFrame(rows)


def detect_leakage_cols(cols: list[str]) -> list[str]:
    def canonical(name: str) -> str:
        s = str(name).strip().lower()
        s = re.sub(r"[^a-z0-9]+", "_", s)
        return re.sub(r"_+", "_", s).strip("_")

    drop = set()
    for raw in cols:
        c = canonical(raw)
        if c in {"id", "row_id", "record_id", "flow_id", "flowid", "flow_identifier", "packet_id", "session_id"}:
            drop.add(c)
        if "timestamp" in c or c in {"time", "datetime"} or c.endswith("_time") or c.endswith("_ts") or c.endswith("_date"):
            drop.add(c)
        if c in {"ip", "srcip", "dstip"} or c.endswith("ip_dec") or re.search(r"(^|_)ip($|_)", c):
            drop.add(c)
        if "ip" in c and any(k in c for k in ["src", "dst", "source", "dest", "destination", "local"]):
            drop.add(c)
        if c in {"port", "sport", "dsport", "srcport", "dstport"} or c.endswith("_port") or re.search(r"(^|_)port($|_)", c):
            drop.add(c)
        if "port" in c and any(k in c for k in ["src", "dst", "source", "dest", "destination", "local"]):
            drop.add(c)
    return sorted(drop)


def first_raw_header(dataset: str) -> list[str]:
    if dataset == "CICIDS2017":
        root = ROOT / "Datasets" / "CICIDS 2017"
        files = sorted(p for p in root.glob("*.csv") if "features" not in p.name.lower() and "_plus" not in p.name.lower())
    else:
        root = ROOT / "Datasets" / "CIC IoT Dataset 2023"
        files = sorted(p for p in root.glob("part-*.csv") if "features" not in p.name.lower())
    if not files:
        return []
    return list(pd.read_csv(files[0], nrows=0).columns)


def support_report_row(dataset: str, split_name: str, report_path: Path, eligible_path: Path) -> dict[str, Any]:
    report = read_json(report_path)
    best = report.get("best_eval", {})
    eligible = pd.read_csv(eligible_path) if eligible_path.exists() else pd.DataFrame()
    raw_cols = first_raw_header(dataset)
    leakage_cols = detect_leakage_cols(raw_cols)
    used_path = (
        CFG["processed_cicids_recovery"] / "B_day_file" / "CICIDS2017" / "USED_COLUMNS.json"
        if dataset == "CICIDS2017"
        else CFG["processed_v5"] / "B_day_file" / "CICIoT2023" / "USED_COLUMNS.json"
    )
    used = read_json(used_path).get("columns", [])
    return {
        "dataset": dataset,
        "split_or_audit_surface": split_name,
        "scenario_valid": bool(best.get("valid", len(eligible) > 0)),
        "eligible_holdouts": int(len(eligible)),
        "eligible_holdout_list": "|".join(sorted(eligible["holdout_family"].astype(str).tolist())) if "holdout_family" in eligible.columns else "",
        "raw_column_count_first_file": int(len(raw_cols)),
        "leakage_like_raw_columns_detected": int(len(leakage_cols)),
        "leakage_like_raw_column_names": "|".join(leakage_cols),
        "processed_column_count": int(len(used)),
        "port_bucket_columns_retained": int(sum(str(c).endswith("_bucket") for c in used)),
        "audit_note": "accepted thesis evidence" if len(eligible) > 0 else "rejected or diagnostic surface",
    }


def build_support_leakage_summary() -> pd.DataFrame:
    rows = [
        support_report_row(
            "CICIDS2017",
            "whole-file Protocol B diagnostic",
            CFG["processed_cicids_recovery"] / "B_day_file" / "CICIDS2017" / "protocol_b_wholefile_search_report.json",
            CFG["processed_cicids_recovery"] / "B_day_file" / "CICIDS2017" / "protocol_b_wholefile_eligible_holdouts.csv",
        ),
        support_report_row(
            "CICIDS2017",
            "recovered contiguous-within-day Protocol B",
            CFG["processed_cicids_recovery"] / "B_day_file" / "CICIDS2017" / "protocol_b_search_report.json",
            CFG["support_audit_cicids_recovery"] / "eligible_holdouts_all.csv",
        ),
        support_report_row(
            "CICIoT2023",
            "day-file Protocol B baseline",
            CFG["processed_v5"] / "B_day_file" / "CICIoT2023" / "protocol_b_search_report.json",
            CFG["processed_v5"] / "B_day_file" / "CICIoT2023" / "protocol_b_eligible_holdouts.csv",
        ),
    ]
    return pd.DataFrame(rows)


def plot_threshold_sensitivity(curves: pd.DataFrame) -> Path:
    selected = set(CFG["selected_curve_holdouts"])
    fig, axes = plt.subplots(2, 2, figsize=(11, 8), sharex=False, sharey=True)
    methods = ["control_tau", "margin_reject", "entropy_reject", "conformal_reject"]
    colors = {
        "control_tau": "#355C7D",
        "margin_reject": "#6C5B7B",
        "entropy_reject": "#C06C84",
        "conformal_reject": "#2A9D8F",
    }
    for ax, (dataset, holdout) in zip(axes.ravel(), CFG["selected_curve_holdouts"]):
        sub = curves.loc[(curves["dataset"] == dataset) & (curves["holdout_family"] == holdout)].copy()
        for method in methods:
            m = sub.loc[sub["method"] == method].sort_values("false_unknown_rate_all_known")
            if m.empty:
                continue
            ax.plot(
                pd.to_numeric(m["false_unknown_rate_all_known"], errors="coerce"),
                pd.to_numeric(m["unknown_detection_rate"], errors="coerce"),
                marker="o",
                linewidth=1.5,
                markersize=3,
                label=method.replace("_", " "),
                color=colors.get(method),
            )
        ax.set_title(f"{dataset} {holdout}", fontsize=10)
        ax.set_xlabel("False-unknown rate")
        ax.set_ylabel("Unknown-detection rate")
        ax.grid(True, alpha=0.25)
    handles, labels = axes.ravel()[0].get_legend_handles_labels()
    fig.legend(handles, labels, loc="lower center", ncol=4, frameon=False)
    fig.suptitle("Threshold Sensitivity For Selected Holdouts", fontsize=13, y=0.98)
    path = CFG["out_root"] / "figures" / "threshold_sensitivity_selected_holdouts.png"
    savefig(path)
    return path


def plot_support_leakage(summary: pd.DataFrame) -> Path:
    labels = [f"{r.dataset}\n{r.split_or_audit_surface.split(' Protocol')[0]}" for r in summary.itertuples()]
    fig, axes = plt.subplots(1, 2, figsize=(11, 4))
    axes[0].bar(labels, summary["eligible_holdouts"], color=["#B56576", "#2A9D8F", "#457B9D"])
    axes[0].set_ylabel("Eligible holdouts")
    axes[0].set_title("Support-Audit Outcome")
    axes[0].tick_params(axis="x", labelrotation=20)
    axes[1].bar(labels, summary["leakage_like_raw_columns_detected"], color=["#B56576", "#2A9D8F", "#457B9D"])
    axes[1].set_ylabel("Leakage-like raw columns detected")
    axes[1].set_title("Leakage-Control Surface")
    axes[1].tick_params(axis="x", labelrotation=20)
    path = CFG["out_root"] / "figures" / "support_leakage_audit_visual.png"
    savefig(path)
    return path


def plot_sink_bridge(sink_diag: pd.DataFrame) -> Path:
    plot_df = sink_diag.sort_values(["dataset", "holdout_family"]).copy()
    labels = [f"{r.dataset}\n{r.holdout_family}->{r.top_sink_family}" for r in plot_df.itertuples()]
    x = np.arange(len(plot_df))
    fig, ax = plt.subplots(figsize=(12, 5))
    ax.bar(x - 0.18, plot_df["control_tau_top_sink_share"], width=0.36, label="Top-sink share", color="#457B9D")
    ax.bar(x + 0.18, plot_df["overconfident_sink_absorption_rate"], width=0.36, label="Overconfident sink absorption", color="#E76F51")
    ax.set_xticks(x)
    ax.set_xticklabels(labels, rotation=45, ha="right", fontsize=8)
    ax.set_ylim(0, 1.05)
    ax.set_ylabel("Rate")
    ax.set_title("Aggregate Sink Collapse Linked To High-Confidence Examples")
    ax.legend(frameon=False)
    ax.grid(True, axis="y", alpha=0.25)
    path = CFG["out_root"] / "figures" / "sink_to_example_bridge.png"
    savefig(path)
    return path


def plot_deployment_feasibility(taxonomy: pd.DataFrame, runtime: pd.DataFrame) -> Path:
    fig, axes = plt.subplots(1, 2, figsize=(12, 4.8))
    counts = taxonomy["difficulty_category"].value_counts().sort_index()
    axes[0].bar(counts.index, counts.values, color="#2A9D8F")
    axes[0].set_title("Holdout Difficulty Categories")
    axes[0].set_ylabel("Holdouts")
    axes[0].tick_params(axis="x", rotation=35)
    cost = runtime.loc[runtime["lane"].isin(["Step 5 open-set baselines", "Step 6 sink-aware rejection", "Drift action study", "Thesis finalization pack"])].copy()
    axes[1].bar(cost["lane"], pd.to_numeric(cost["total_size_mb"], errors="coerce"), color="#6C5B7B")
    axes[1].set_title("Artifact Footprint Of Deployment-Relevant Layers")
    axes[1].set_ylabel("Artifact size (MB)")
    axes[1].tick_params(axis="x", rotation=35)
    path = CFG["out_root"] / "figures" / "reliability_deployment_feasibility.png"
    savefig(path)
    return path


def build_manifest() -> pd.DataFrame:
    rows: list[dict[str, Any]] = []
    for item_id, filename, description in TABLES:
        path = CFG["out_root"] / "tables" / filename
        row_count = len(pd.read_csv(path)) if path.exists() else 0
        rows.append(
            {
                "item_id": item_id,
                "item_type": "table",
                "description": description,
                "path": str(path.relative_to(ROOT)),
                "exists": path.exists(),
                "non_empty": bool(path.exists() and row_count > 0),
                "row_count": row_count,
            }
        )
    for item_id, filename, description in FIGURES:
        path = CFG["out_root"] / "figures" / filename
        rows.append(
            {
                "item_id": item_id,
                "item_type": "figure",
                "description": description,
                "path": str(path.relative_to(ROOT)),
                "exists": path.exists(),
                "non_empty": bool(path.exists() and path.stat().st_size > 0),
                "row_count": "",
            }
        )
    md_path = CFG["manuscript_dir"] / str(CFG["final_md_name"])
    docx_path = CFG["manuscript_dir"] / str(CFG["final_docx_name"])
    for item_id, path, desc in [
        ("M20", md_path, "Final v3 manuscript markdown"),
        ("M21", docx_path, "Final v3 manuscript docx"),
    ]:
        rows.append(
            {
                "item_id": item_id,
                "item_type": "manuscript",
                "description": desc,
                "path": str(path.relative_to(ROOT)),
                "exists": path.exists(),
                "non_empty": bool(path.exists() and path.stat().st_size > 0),
                "row_count": "",
            }
        )
    manifest = pd.DataFrame(rows)
    manifest.to_csv(CFG["out_root"] / "output_manifest.csv", index=False)
    return manifest


def fmt(value: Any) -> str:
    try:
        numeric = float(value)
    except (TypeError, ValueError):
        return "" if value is None else str(value)
    if math.isnan(numeric):
        return ""
    if abs(numeric) >= 1000:
        return f"{numeric:,.0f}"
    return f"{numeric:.4f}"


def markdown_table(df: pd.DataFrame, cols: list[tuple[str, str]], max_rows: int | None = None) -> str:
    use = df.copy()
    if max_rows is not None:
        use = use.head(max_rows)
    header = "| " + " | ".join(label for _, label in cols) + " |"
    sep = "| " + " | ".join("---" for _ in cols) + " |"
    lines = [header, sep]
    for rec in use.to_dict("records"):
        lines.append("| " + " | ".join(fmt(rec.get(col, "")) for col, _ in cols) + " |")
    return "\n".join(lines)


def update_manuscript_markdown(
    drift_summary: pd.DataFrame,
    taxonomy: pd.DataFrame,
    sink_diag: pd.DataFrame,
    selected_stats: pd.DataFrame,
    runtime: pd.DataFrame,
) -> Path:
    src = CFG["source_manuscript"]
    text = src.read_text(encoding="utf-8")
    text = text.replace("Complete manuscript v1 - 2026-05-18", "Final manuscript v3 - 2026-06-02")
    text = text.replace(
        "# Chapter 7. Conclusion and Thesis Completion Plan",
        "# Chapter 7. Conclusion and Final Submission Boundary",
    )

    hardening = "\n".join(
        [
            "## 4.11 Final Reliability And Deployment Hardening",
            "",
            "The final thesis hardening layer does not introduce new classifier training. It promotes already-materialized diagnostics into the main results so that reliability, rejection cost, calibration, and deployment feasibility are visible in the defended narrative.",
            "",
            "**Table 4.8. Cross-dataset drift-action summary.**",
            "",
            markdown_table(
                drift_summary,
                [
                    ("dataset", "Dataset"),
                    ("trigger_source", "Trigger"),
                    ("n_reviewed_windows", "Reviewed"),
                    ("n_triggered_windows", "Triggered"),
                    ("mean_best_delta_macro_f1_vs_static", "Mean delta macro-F1"),
                    ("mean_best_delta_system_ece_vs_static", "Mean delta ECE"),
                    ("most_common_best_action", "Most common action"),
                ],
                max_rows=8,
            ),
            "",
            "**Table 4.9. Holdout difficulty taxonomy.**",
            "",
            markdown_table(
                taxonomy,
                [
                    ("dataset", "Dataset"),
                    ("holdout_family", "Holdout"),
                    ("difficulty_category", "Category"),
                    ("unknown_detection_rate", "Unknown detection"),
                    ("top_sink_family", "Top sink"),
                    ("open_set_review_decision", "Rejector decision"),
                ],
            ),
            "",
            "**Table 4.10. Selected statistical results moved into the main body.**",
            "",
            markdown_table(
                selected_stats,
                [
                    ("result_family", "Result family"),
                    ("dataset", "Dataset"),
                    ("holdout_family", "Holdout"),
                    ("method_or_lane", "Method/lane"),
                    ("decision", "Decision"),
                    ("delta_macro_f1", "Delta macro-F1"),
                    ("delta_unknown_detection", "Delta unknown detection"),
                ],
            ),
            "",
            "![Threshold-sensitivity curves for selected holdouts.](thesis_finalization_pack/figures/threshold_sensitivity_selected_holdouts.png)",
            "",
            "![Support and leakage audit visual.](thesis_finalization_pack/figures/support_leakage_audit_visual.png)",
            "",
            "![Sink-to-example bridge.](thesis_finalization_pack/figures/sink_to_example_bridge.png)",
            "",
        ]
    )
    text = text.replace("# Chapter 5. Discussion", hardening + "\n# Chapter 5. Discussion")

    discussion_insert = "\n".join(
        [
            "## 5.9 Calibration, Error Cases, And Deployment Feasibility",
            "",
            "The final calibration diagnostics show that sink collapses are not only aggregate confusion-matrix events. They often involve confident absorption of a withheld family into a nearby known family, which is operationally different from harmless uncertainty. Representative error cases in `thesis_finalization_pack/tables/T27_representative_error_cases.csv` connect row-level confidence, margin, entropy, and predicted sink family back to the aggregate sink-collapse table.",
            "",
            "The deployment interpretation is therefore bounded. Strong and mixed holdouts can support ordinary monitored operation; hard-collapse cases require explicit escalation or rejection policy; targeted-recovery cases justify a narrow intervention; and negative-tradeoff cases warn against deploying a rejector only because it raises unknown detection. Runtime/resource reporting is also deliberately conservative: the artifact roots record worker settings and storage footprint, but exact historical wall-clock timing is only available where execution settings or logs preserved it.",
            "",
            "![Reliability and deployment feasibility summary.](thesis_finalization_pack/figures/reliability_deployment_feasibility.png)",
            "",
        ]
    )
    text = text.replace("# Chapter 6. Limitations", discussion_insert + "# Chapter 6. Limitations")

    new_ch7 = "\n".join(
        [
            "# Chapter 7. Conclusion and Final Submission Boundary",
            "",
            "This thesis shows that support-audited trustworthiness evaluation changes how two-stage ML-based IDS should be interpreted. The strict closed-set baseline is strong enough to anchor the work, but Protocol B reveals that unknown-family behavior is uneven, split-dependent, and structured by sink-family collapses. Calibration-aware abstention and alternative rejectors improve selected cases, but no universal rejector is supported. Drift and explanation stability add reliability context, while external robustness results warn against unsupported transfer claims.",
            "",
            "The finalization pass completes the bounded thesis-level refinements identified after the seminar presentation. Drift-to-action is now evaluated across both primary datasets and is reported as a dataset-conditional maintenance result. Per-family and sink-family calibration diagnostics show where confident absorption rather than mere low confidence drives failure. Threshold-sensitivity curves, selected confidence intervals, holdout taxonomy, support/leakage-audit visuals, representative error cases, and runtime/resource-cost summaries are now part of the final thesis evidence package.",
            "",
            "The final submission boundary remains disciplined. The audited open-set benchmark is promoted as a contribution because it exposes family-dependent behavior and method heterogeneity. Sink-aware rejection remains a targeted case study, not a broad replacement method. External datasets remain robustness stress tests, not transfer validation. Larger directions such as Bayesian family-aware abstention, temporal stochastic modeling, deep learning baselines, and validated cross-dataset transfer are future work rather than missing thesis requirements.",
            "",
        ]
    )
    text = re.sub(
        r"# Chapter 7\. Conclusion.*?(?=# References)",
        new_ch7,
        text,
        flags=re.S,
    )

    appendix_insert = "\n".join(
        [
            "# Appendix E. Finalization Pack Tables",
            "",
            "The finalization pack tables are generated by `20.BuildThesisFinalizationPack.py` and stored under `thesis_finalization_pack/tables/`. The full confidence-interval and paired-test appendices remain in `paper1_benchmark_pack/appendix/`; only selected results are moved into the main body.",
            "",
            "**Appendix Table E.1. Sink calibration diagnostic excerpt.**",
            "",
            markdown_table(
                sink_diag,
                [
                    ("dataset", "Dataset"),
                    ("holdout_family", "Holdout"),
                    ("top_sink_family", "Sink"),
                    ("control_tau_top_sink_share", "Sink share"),
                    ("mean_sink_confidence", "Mean sink confidence"),
                    ("overconfident_sink_absorption_rate", "Overconfident sink absorption"),
                    ("calibration_failure_signal", "Signal"),
                ],
                max_rows=12,
            ),
            "",
            "**Appendix Table E.2. Runtime/resource-cost summary.**",
            "",
            markdown_table(
                runtime,
                [
                    ("lane", "Lane"),
                    ("file_count", "Files"),
                    ("total_size_mb", "Size MB"),
                    ("case_parallel_workers", "Workers"),
                    ("threads_per_worker", "Threads"),
                    ("wallclock_evidence_status", "Wall-clock evidence"),
                ],
            ),
            "",
        ]
    )
    text = text.replace("# Appendix A. Artifact Map", appendix_insert + "\n# Appendix A. Artifact Map")

    out = CFG["manuscript_dir"] / str(CFG["final_md_name"])
    out.write_text(text, encoding="utf-8")
    return out


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: object, *, header: bool = False, font_size: float = 7.5) -> None:
    cell.text = "" if text is None else str(text)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        for run in paragraph.runs:
            run.font.size = Pt(font_size)
            if header:
                run.bold = True
    if header:
        shade_cell(cell, "D9EAF7")


def add_inline_runs(paragraph, text: str) -> None:
    for part in INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Aptos Mono"
            run._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Mono")
            run._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Mono")
        else:
            paragraph.add_run(part)


def add_table(doc: Document, rows: list[list[str]], *, font_size: float = 7.5) -> None:
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row_idx, row in enumerate(rows):
        if row_idx == 0:
            tr_pr = table.rows[row_idx]._tr.get_or_add_trPr()
            tbl_header = OxmlElement("w:tblHeader")
            tbl_header.set(qn("w:val"), "true")
            tr_pr.append(tbl_header)
        for col_idx in range(col_count):
            value = row[col_idx] if col_idx < len(row) else ""
            set_cell_text(table.cell(row_idx, col_idx), value, header=row_idx == 0, font_size=font_size)
    doc.add_paragraph()


def add_image(doc: Document, path_text: str, alt: str = "") -> None:
    path = (ROOT / path_text).resolve()
    if not path.exists():
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Inches(6.4))
    if alt:
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = caption.add_run(alt)
        run.italic = True
        run.font.size = Pt(8.5)


def parse_front_matter(lines: list[str]) -> tuple[dict[str, str], list[str]]:
    meta: dict[str, str] = {}
    body_start = 0
    if lines and lines[0].strip() == "---":
        for idx in range(1, len(lines)):
            line = lines[idx]
            if line.strip() == "---":
                body_start = idx + 1
                break
            if ":" in line:
                key, value = line.split(":", 1)
                meta[key.strip()] = value.strip().strip('"')
    return meta, lines[body_start:]


def parse_markdown_into_doc(doc: Document, body: list[str]) -> None:
    paragraph_buffer: list[str] = []
    idx = 0
    while idx < len(body):
        raw = body[idx]
        stripped = raw.strip()
        if not stripped:
            if paragraph_buffer:
                paragraph = doc.add_paragraph()
                add_inline_runs(paragraph, " ".join(paragraph_buffer))
                paragraph_buffer = []
            idx += 1
            continue
        if stripped.startswith("|"):
            if paragraph_buffer:
                paragraph = doc.add_paragraph()
                add_inline_runs(paragraph, " ".join(paragraph_buffer))
                paragraph_buffer = []
            table_lines: list[str] = []
            while idx < len(body) and body[idx].strip().startswith("|"):
                table_lines.append(body[idx])
                idx += 1
            rows: list[list[str]] = []
            for line in table_lines:
                cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
                if all(set(cell) <= set("-: ") for cell in cells):
                    continue
                rows.append(cells)
            add_table(doc, rows)
            continue
        image_match = re.match(r"!\[(.*?)\]\((.*?)\)", stripped)
        if image_match:
            if paragraph_buffer:
                paragraph = doc.add_paragraph()
                add_inline_runs(paragraph, " ".join(paragraph_buffer))
                paragraph_buffer = []
            add_image(doc, image_match.group(2), image_match.group(1))
            idx += 1
            continue
        heading_match = re.match(r"^(#{1,3})\s+(.*)$", stripped)
        if heading_match:
            if paragraph_buffer:
                paragraph = doc.add_paragraph()
                add_inline_runs(paragraph, " ".join(paragraph_buffer))
                paragraph_buffer = []
            level = len(heading_match.group(1))
            title = heading_match.group(2)
            if level == 1 and (title.startswith("Chapter ") or title.startswith("References") or title.startswith("Appendix ")):
                doc.add_page_break()
            paragraph = doc.add_heading(level=level)
            add_inline_runs(paragraph, title)
            idx += 1
            continue
        if re.match(r"^\d+\.\s+", stripped):
            if paragraph_buffer:
                paragraph = doc.add_paragraph()
                add_inline_runs(paragraph, " ".join(paragraph_buffer))
                paragraph_buffer = []
            paragraph = doc.add_paragraph(style="List Number")
            add_inline_runs(paragraph, re.sub(r"^\d+\.\s+", "", stripped))
            idx += 1
            continue
        if stripped.startswith("- "):
            if paragraph_buffer:
                paragraph = doc.add_paragraph()
                add_inline_runs(paragraph, " ".join(paragraph_buffer))
                paragraph_buffer = []
            paragraph = doc.add_paragraph(style="List Bullet")
            add_inline_runs(paragraph, stripped[2:])
            idx += 1
            continue
        paragraph_buffer.append(stripped)
        idx += 1
    if paragraph_buffer:
        paragraph = doc.add_paragraph()
        add_inline_runs(paragraph, " ".join(paragraph_buffer))


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    styles["Normal"]._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    styles["Normal"].font.size = Pt(10.5)
    for name, size, color in [
        ("Title", 19, "1F4E79"),
        ("Heading 1", 15, "1F4E79"),
        ("Heading 2", 12.5, "2F5597"),
        ("Heading 3", 11.5, "385723"),
    ]:
        style = styles[name]
        style.font.name = "Aptos Display" if name != "Heading 3" else "Aptos"
        style._element.rPr.rFonts.set(qn("w:ascii"), style.font.name)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), style.font.name)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True


def build_docx(md_path: Path) -> Path:
    meta, body = parse_front_matter(md_path.read_text(encoding="utf-8").splitlines())
    doc = Document()
    configure_styles(doc)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(meta.get("title", "Thesis Manuscript"))
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(meta.get("author", ""))
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(11)
    status = doc.add_paragraph()
    status.alignment = WD_ALIGN_PARAGRAPH.CENTER
    status.add_run(meta.get("date", ""))
    doc.add_page_break()
    parse_markdown_into_doc(doc, body)
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        if not footer.text:
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = footer.add_run("Final v3 thesis manuscript - generated from repo evidence package")
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(100, 100, 100)
    out = CFG["manuscript_dir"] / str(CFG["final_docx_name"])
    doc.save(out)
    return out


def main() -> None:
    mkdirs()
    b_best = read_csv("protocol_b_best_per_holdout_standardized.csv")
    failure = read_csv("protocol_b_failure_modes.csv")
    notes = read_csv("protocol_b_holdout_notes.csv")
    open_recs = read_csv("open_set_holdout_recommendations.csv")
    sink_recs = read_csv("sink_aware_holdout_recommendations.csv")
    curves = read_csv("open_set_unknown_known_curves.csv")

    drift_summary_path = CFG["drift_action_root"] / "action_policy_summary.csv"
    drift_dataset_path = CFG["drift_action_root"] / "action_policy_dataset_summary.csv"
    drift_summary = pd.read_csv(drift_summary_path) if drift_summary_path.exists() else pd.DataFrame()
    drift_dataset = pd.read_csv(drift_dataset_path) if drift_dataset_path.exists() else pd.DataFrame()
    if not drift_dataset.empty:
        drift_table = drift_summary.merge(
            drift_dataset[
                [
                    "dataset",
                    "n_stable_windows",
                    "n_drift_windows",
                    "mean_best_delta_macro_f1_vs_static_all_windows",
                    "interpretation",
                ]
            ],
            on="dataset",
            how="left",
        )
    else:
        drift_table = drift_summary.copy()

    per_family = build_per_family_calibration(b_best)
    sink_diag, examples = build_sink_calibration_and_examples(failure, per_family)
    taxonomy = build_holdout_taxonomy(b_best, notes, open_recs, sink_recs, sink_diag)
    tradeoff = build_rejection_tradeoff(open_recs)
    runtime = build_runtime_summary()
    selected_stats = build_selected_statistics(open_recs, sink_recs)
    support_summary = build_support_leakage_summary()

    write_table(drift_table, "drift_action_cross_dataset_summary.csv")
    write_table(per_family, "per_family_calibration.csv")
    write_table(sink_diag, "sink_calibration_diagnostic.csv")
    write_table(taxonomy, "holdout_difficulty_taxonomy.csv")
    write_table(tradeoff, "rejection_tradeoff_summary.csv")
    write_table(runtime, "runtime_resource_cost_summary.csv")
    write_table(selected_stats, "selected_statistical_summary.csv")
    write_table(examples, "representative_error_cases.csv")
    write_table(support_summary, "support_leakage_audit_summary.csv")

    plot_threshold_sensitivity(curves)
    plot_support_leakage(support_summary)
    plot_sink_bridge(sink_diag)
    runtime = build_runtime_summary()
    write_table(runtime, "runtime_resource_cost_summary.csv")
    plot_deployment_feasibility(taxonomy, runtime)

    md_path = update_manuscript_markdown(drift_table, taxonomy, sink_diag, selected_stats, runtime)
    build_docx(md_path)
    manifest = build_manifest()
    print(f"Wrote thesis finalization pack to: {CFG['out_root']}")
    print(f"Manifest rows: {len(manifest)}")


if __name__ == "__main__":
    main()
